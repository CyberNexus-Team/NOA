from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from io import BytesIO
from docx import Document
from docx.shared import Pt

app = Flask(__name__)
CORS(app)

@app.get("/healthz")
def healthz():
    return "ok\n", 200

@app.post("/convert")
def convert():
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "missing file field 'file'"}), 400
    html = f.read().decode("utf-8", errors="replace")

    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run(html)
    run.font.size = Pt(10)

    out = BytesIO()
    doc.save(out); out.seek(0)
    filename = request.form.get("filename") or "report.docx"
    return send_file(
        out,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename,
    )
